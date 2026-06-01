from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/Unio_Rentals_Platform_Gap_SOP_Report.docx"


def set_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell(cell, header, bold=True)
        shading = cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fill = OxmlElement("w:shd")
        fill.set(qn("w:fill"), "E8EEF5")
        shading.append(fill)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Inches(widths[index])
            set_cell(cells[index], value)

    document.add_paragraph()
    return table


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_steps(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_flow(document, title, lines):
    document.add_heading(title, level=2)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(31, 77, 120)


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor(11, 37, 69)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    document.add_paragraph("Unio Rentals", style="Title")
    subtitle = document.add_paragraph(
        "Rental service provider platform gap analysis, SOP, workflow map, and release readiness brief"
    )
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    document.add_paragraph("Prepared after code review and implementation pass | May 2026")

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "Unio Rentals now has the foundation of a marketplace: browsing, listing publication, login, OTP "
        "registration, quick booking, dashboard, booking history, and admin statistics. The latest implementation "
        "pass tightened route protection, reduced client image payload risk, added authenticated property publishing, "
        "added a user-owned listing view, and reduced initial frontend workload by removing eager route preloading."
    )
    document.add_paragraph(
        "The platform is not yet production-complete. The largest remaining gaps are persistent property and booking "
        "storage, production-grade image hosting, real payment integration, admin moderation, landlord KYC, refund and "
        "cancellation flows, notification workflows, and cookie-based auth with refresh-token rotation."
    )

    document.add_heading("2. Completed In This Development Pass", level=1)
    add_bullets(
        document,
        [
            "Protected property publishing and booking endpoints behind authentication.",
            "Added /api/properties/mine so users can access properties listed from their own account.",
            "Added dashboard section for the user's listed properties.",
            "Added typed location search in the top navigation and home booking bar.",
            "Added backend-backed location suggestions and location-aware property filtering.",
            "Connected current-location permission to the shared selected location state.",
            "Compressed uploaded property/add-on images in the browser before submission to reduce UI freezes and payload size.",
            "Moved frontend access token storage from persistent localStorage to sessionStorage and scrubbed legacy localStorage values on new login/logout.",
            "Added basic backend security headers and disabled the Express powered-by header.",
            "Guarded account-only frontend routes such as publish property, orders, coins, and settings.",
            "Removed eager preloading of all lazy routes to improve first-load smoothness.",
            "Added autocomplete attributes for login, signup, admin login, and OTP fields.",
        ],
    )

    document.add_heading("3. Endpoint Availability Matrix", level=1)
    add_table(
        document,
        ["Area", "Endpoint", "Current Status", "Production Remark"],
        [
            ["Health", "GET /health", "Available", "Add uptime monitoring and deployment health checks."],
            ["Auth", "POST /api/auth/register/request-otp", "Available", "Hash OTP before storage and add email delivery observability."],
            ["Auth", "POST /api/auth/register/verify-otp", "Available", "Add OTP replay protection and audit trail."],
            ["Auth", "POST /api/auth/login", "Available", "Move final production design to httpOnly cookies plus refresh rotation."],
            ["Auth", "GET /api/auth/me", "Available", "Works with bearer token; should support cookie session later."],
            ["Properties", "GET /api/properties", "Available", "Currently backed by in-memory marketplace data."],
            ["Locations", "GET /api/locations/suggest", "Available", "Suggests known city/locality values and nearest known location from coordinates."],
            ["Properties", "GET /api/properties/:id", "Available", "Needs persistent database lookup and public/private field shaping."],
            ["Properties", "POST /api/properties", "Protected", "Now authenticated; still needs admin approval and object storage images."],
            ["Properties", "GET /api/properties/mine", "Protected", "Now available for user-owned listings."],
            ["Bookings", "GET /api/bookings", "Protected", "Now scoped by logged-in user; needs database persistence."],
            ["Bookings", "POST /api/bookings", "Protected", "Needs real payment intent, idempotency, and concurrency lock."],
            ["Admin", "GET /api/admin/dashboard", "Protected", "Stats exist; needs moderation queues and operational views."],
        ],
        [1.1, 2.15, 1.15, 2.1],
    )

    document.add_heading("4. Missing Platform Capabilities", level=1)
    add_table(
        document,
        ["Priority", "Missing Capability", "Why It Matters", "Recommended Action"],
        [
            ["Critical", "Persistent property and booking APIs", "In-memory data is lost on restart and cannot support many users.", "Use Prisma models for properties, photos, amenities, bookings, payments, and invoices."],
            ["Critical", "Production image storage", "Base64 images in JSON are slow, expensive, and unsafe at scale.", "Upload to S3/Cloudinary/Azure Blob through signed URLs; store only image URLs."],
            ["Critical", "Payment gateway and idempotency", "Current payment is simulated as paid.", "Integrate Razorpay/Stripe; use payment intents, webhook verification, and idempotency keys."],
            ["Critical", "Admin property moderation", "User listings go live immediately.", "Add pending approval queue, reject reasons, and listing status transitions."],
            ["High", "Cookie-based auth hardening", "Bearer tokens in browser storage are exposed to XSS.", "Use httpOnly Secure SameSite cookies, CSRF protection, refresh rotation, and token revocation."],
            ["High", "Landlord onboarding/KYC", "Trust is central in rental marketplaces.", "Collect document verification, address proof, bank details, and approval status."],
            ["High", "Search and filters", "Users need price, location, type, duration, amenities, availability, and map filters.", "Move search to backend query parameters with indexes."],
            ["High", "Geo coverage expansion", "Current location suggestions use known locations plus listed property locations.", "Add production geocoding/autocomplete provider, serviceable-area management, and map radius search."],
            ["Medium", "Cancellation and refund workflows", "Bookings need lifecycle management after payment.", "Add cancellation policy, refund status, and admin override flow."],
            ["Medium", "Notifications", "Users need OTP, booking, payment, listing approval, and support updates.", "Add email/SMS/WhatsApp notification service with templates and retry tracking."],
            ["Medium", "Observability and load protection", "Multiple users require visibility into errors, latency, and traffic.", "Add structured logs, request IDs, metrics, alerts, CDN, and cache strategy."],
        ],
        [0.8, 1.6, 2.0, 2.1],
    )

    document.add_heading("5. Workflow Charts", level=1)
    add_flow(
        document,
        "Tenant Rental Flow",
        [
            "Browse rentals -> Filter/search -> View listing detail",
            "       -> Login/Register -> Select dates -> Create payment intent",
            "       -> Pay -> Confirm booking -> Invoice + landlord contact",
            "       -> My Orders -> Support/cancel/refund if needed",
        ],
    )
    add_flow(
        document,
        "Landlord Listing Flow",
        [
            "Login/Register -> Complete landlord profile/KYC -> Publish property draft",
            "       -> Upload photos to object storage -> Submit for admin review",
            "       -> Admin approves/rejects -> Listing becomes searchable",
            "       -> Booking request/payment -> Calendar blocked -> Payout after check-in",
        ],
    )
    add_flow(
        document,
        "Admin Operations Flow",
        [
            "Admin login -> Dashboard -> Review pending properties",
            "       -> Verify KYC/photos/pricing/rules -> Approve or reject",
            "       -> Monitor bookings/payments/refunds -> Resolve support cases",
            "       -> Review reports/fraud signals -> Suspend bad listings/users",
        ],
    )

    document.add_heading("6. SOP: Property Listing", level=1)
    add_steps(
        document,
        [
            "Landlord logs in and completes profile information.",
            "Landlord enters title, location, property type, rent, rules, facilities, landmark, and food options.",
            "Landlord uploads minimum 3 clear property photos. The frontend compresses images; production should upload originals to object storage.",
            "System saves listing as PENDING_APPROVAL and notifies admin.",
            "Admin checks photo quality, duplicate listing risk, pricing, contact validity, and policy compliance.",
            "Admin approves listing. The listing becomes AVAILABLE and searchable.",
            "Landlord can view all own listings from the dashboard and update availability or pricing after approval workflow.",
        ],
    )

    document.add_heading("7. SOP: Booking And Payment", level=1)
    add_steps(
        document,
        [
            "Tenant logs in before booking.",
            "Tenant selects check-in/check-out dates and reviews calculated rent.",
            "Backend validates date availability using a transaction or row lock.",
            "Backend creates payment intent with an idempotency key.",
            "Payment gateway confirms payment through a signed webhook.",
            "Backend marks booking CONFIRMED, payment PAID, generates invoice, and blocks the property calendar.",
            "Tenant sees the booking in My Orders; landlord receives booking notification.",
        ],
    )

    document.add_heading("8. SOP: Security And Support", level=1)
    document.add_heading("8A. SOP: Location Search And Discovery", level=2)
    add_steps(
        document,
        [
            "User types a city/locality in the top location field or chooses current location.",
            "Frontend calls /api/locations/suggest and displays matching serviceable locations.",
            "User selects a suggestion or searches the typed location.",
            "Frontend stores the selected location and requests /api/properties with the location query.",
            "Backend filters listings by city/locality tokens and returns user-published properties first when they match.",
            "Home and quick-rent pages render results, or show an empty state with a prompt to list the first property there.",
        ],
    )
    add_bullets(
        document,
        [
            "Never log passwords, OTPs, raw tokens, payment secrets, or identity documents.",
            "Use environment-managed secrets and rotate JWT/payment/email keys regularly.",
            "Use request IDs for support investigations and payment reconciliation.",
            "Run dependency audit, TypeScript build, and smoke tests before every release.",
            "Create admin escalation policy for disputed booking, fake listing, payment failure, and refund requests.",
        ],
    )

    document.add_heading("9. Release Readiness Roadmap", level=1)
    add_table(
        document,
        ["Phase", "Goal", "Key Deliverables", "Exit Criteria"],
        [
            ["Phase 1", "MVP stability", "Database-backed properties/bookings, image URLs, auth guards, dashboard listing ownership.", "No in-memory core marketplace data in production."],
            ["Phase 2", "Trust and payments", "KYC, admin approval, payment gateway, invoices, notifications.", "Bookings require verified payment and approved listing."],
            ["Phase 3", "Marketplace quality", "Search, filters, reviews, cancellation/refund, support desk.", "Users can complete common rental lifecycle without manual ops."],
            ["Phase 4", "Scale", "CDN, caching, queues, observability, load tests, fraud tooling.", "Platform meets target concurrency and reliability SLAs."],
        ],
        [0.85, 1.3, 2.6, 1.65],
    )

    document.add_heading("10. Acceptance Checklist", level=1)
    add_bullets(
        document,
        [
            "Frontend and backend builds pass.",
            "Unauthenticated users cannot publish a property or view account-only routes.",
            "Authenticated user can publish a listing and see it in dashboard under My listed properties.",
            "Uploaded images are compressed client-side and do not freeze the page during normal use.",
            "Booking creation requires login and stores booking against the logged-in user.",
            "Admin dashboard remains protected by admin role.",
            "Production release is blocked until persistent database, object storage, payment gateway, and cookie auth are implemented.",
        ],
    )

    section = document.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Unio Rentals platform brief")

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
